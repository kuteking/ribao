

from docx.enum.style import WD_STYLE_TYPE
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
import time,datetime
def word_doc():
    nt = datetime.datetime.now()
    qnt = datetime.datetime.now()+datetime.timedelta(days=-1)
    
    document = Document()
    document.styles['Normal'].font.name = u'宋体'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    p=document.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r =p.add_run(qnt.strftime('%Y{y}%m{m}%d{d}').format(y='年', m='月', d='日')+'-'+nt.strftime('%Y{y}%m{m}%d{d}').format(y='年', m='月', d='日'))
    r.font.size = Pt(8)
    r.bold = True
    
    table = document.add_table(rows = 3,cols = 3,style = 'Table Grid')#创建一个3x3的表格
    table.autofit = False
    table.cell(0,0).merge(table.cell(1,0))#合并第一列1，2格  合并第三列 1，2格
    table.cell(0,2).merge(table.cell(1,2))

    hdr_cells0 = table.rows[0].cells#用table的row方法可以得到一个表格的一行list其中包含了这一行的所有cell
    note0=hdr_cells0[0].add_paragraph()
    note0.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
    note01=note0.add_run('生产动态')
    note01.bool=True
    note01.font.size = Pt(8) 

    note1=hdr_cells0[1].add_paragraph()
    note1.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
    note11=note1.add_run('工作情况')
    note11.bool=True
    note11.font.size = Pt(8) 

    note2=hdr_cells0[2].add_paragraph()
    note2.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
    note12=note2.add_run('人员及动迁情况')
    note12.bool=True
    note12.font.size = Pt(8) 


    #hdr_cells0[1].add_paragraph('工 作 情 况').alignment=WD_ALIGN_PARAGRAPH.CENTER
    #hdr_cells0[2].add_paragraph('人员及动迁情况').alignment=WD_ALIGN_PARAGRAPH.CENTER

    hdr_cells1 = table.rows[1].cells
    note3=hdr_cells1[1].add_paragraph()
    note3.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
    note13=note3.add_run('当日重点工作及明日工作计划')
    note13.bool=True
    note13.font.size = Pt(8)

    document.save('demo.docx')

    document1 = Document('f:\\08.docx')
    tables=document1.tables[0]
    tables.cell(2,0).text=''
    tables.cell(2,0).text='aaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaa'
    tables.cell(2,1).text=''

    document1.save('demo1.docx')


if __name__=='__main__':
    word_doc()