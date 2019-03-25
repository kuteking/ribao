from tkinter import *
import tkinter.filedialog
import os
import os.path
import tkinter.messagebox
from docx.enum.style import WD_STYLE_TYPE
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
import time,datetime
import open_word
###################
def mywin():
    global zhuant
    global ren_name
    global xj_list
    global qj_list
    global time_data
    def zt_cl():#设备选择
        global zhuant
        global sbtime
        if h1.get()==True: # 因为h1创建时是bool的运算，选中为真，不选为假
            zhuant[0]=1
            sbtime[0]=strh1.get()
        if h2.get() == True:
            zhuant[1]=1
            sbtime[1]=strh2.get()
        if h3.get() == True:
            zhuant[2]=1
            sbtime[2]=strh3.get()
        if h4.get() == True:
            zhuant[3]=1
            sbtime[3]=strh4.get()
        if h5.get() == True:
            zhuant[4]=1
            sbtime[4]=strh5.get()
        if h6.get() == True:
            zhuant[5]=1
            sbtime[5]=strh6.get()
        if h7.get() == True:
            zhuant[6]=1
            sbtime[6]=strh7.get()
        if h8.get() == True:
            zhuant[7]=1
            sbtime[7]=strh8.get()
        if h9.get() == True:
            zhuant[8]=1
            sbtime[8]=strh9.get()
        if h10.get() == True:
            zhuant[9]=1
            sbtime[9]=strh10.get()
        if h11.get() == True:
            zhuant[10]=1
            sbtime[10]=strh11.get()
###############################
    def rs_xz():
        global renshu
        if rs1.get() == True:
            renshu[0]=1
        else:
            renshu[0]=0
        if rs2.get() == True:
            renshu[1]=1
        else:
            renshu[1]=0
        if rs3.get() == True:
            renshu[2]=1
        else:
            renshu[2]=0
        if rs4.get() == True:
            renshu[3]=1
        else:
            renshu[3]=0
        if rs5.get() == True:
            renshu[4]=1
        else:
            renshu[4]=0
        if rs6.get() == True:
            renshu[5]=1
        else:
            renshu[5]=0
##########################
    def xj_xz():
            global xj_list
            if xj1.get() == True:
                xj_list[0]=1
            else:
                xj_list[0]=0
            if xj2.get() == True:
                xj_list[1]=1
            else:
                xj_list[1]=0
            if xj3.get() == True:
                xj_list[2]=1
            else:
                xj_list[2]=0
            if xj4.get() == True:
                xj_list[3]=1
            else:
                xj_list[3]=0
            if xj5.get() == True:
                xj_list[4]=1
            else:
                xj_list[4]=0
            if xj6.get() == True:
                xj_list[5]=1
            else:
                xj_list[5]=0
#############
    def qj_xz():
            global qj_list
            if qj1.get() == True:
                qj_list[0]=1
            else:
                qj_list[0]=0
            if qj2.get() == True:
                qj_list[1]=1
            else:
                qj_list[1]=0
            if qj3.get() == True:
                qj_list[2]=1
            else:
                qj_list[2]=0
            if qj4.get() == True:
                qj_list[3]=1
            else:
                qj_list[3]=0
            if qj5.get() == True:
                qj_list[4]=1
            else:
                qj_list[4]=0
            if qj6.get() == True:
                qj_list[5]=1
            else:
                qj_list[5]=0
########################################
    def bltext():#刷新数据
        global sbtime
        global renshu
        global ren_name
        global xj_list
        global qj_list
        global liulian

        liulian=(llint1.get()+llint1.get()+llint3.get()+llint4.get())*24
        for nn,rena in zip(xj_list,ren_name):
            if nn == 1:
                pass
               
        for nn,rena in zip(qj_list,ren_name):
            if nn == 1:
                pass

        
        ls=0
        for n in renshu:
            if n==1:
                ls+=1
            ren_shu.set(ls)
        sbtime[0]=strh1.get()
        sbtime[1]=strh2.get()
        sbtime[2]=strh3.get()
        sbtime[3]=strh4.get()
        sbtime[4]=strh5.get()
        sbtime[5]=strh6.get()
        sbtime[6]=strh7.get()
        sbtime[7]=strh8.get()
        sbtime[8]=strh9.get()
        sbtime[9]=strh10.get()
        sbtime[10]=strh11.get()
       

        root.after(1000,bltext)
 #################################
    def renyuan():#返回人员情况的字符串
        global sbtime
        global renshu
        global ren_name
        global xj_list
        global qj_list
        strtemp = '潜江站在岗人员：'
        rs=0
        for n in renshu:
            if n==1:
                rs+=1
        strtemp=strtemp+str(rs)+'人\n'

        for nn,rena in zip(renshu,ren_name):
            if nn == 1:
                strtemp=strtemp+''.join(rena)+','
        strtemp=strtemp+'\n'+'夜班人员：1人'+'\n'
        strtemp=strtemp+''.join(ren_name[yb_name.get()])+'\n'

        rs=0
        ls=''
        for nnn,rena1 in zip(xj_list,ren_name):
            if nnn == 1:
                rs=+1
                ls=ls+''.join(rena1)+','
                #print(rena)
        strtemp=strtemp+'休假人员：'+str(rs)+'人\n'
        strtemp=strtemp+ls+'\n'
        rs=0
        ls=''
        for nnnn,rena11 in zip(qj_list,ren_name):
            if nnnn == 1:
                rs+=1
                ls=ls+','.join(renna11)+','
        strtemp=strtemp+'请假人员：'+str(rs)+'人\n'
        strtemp=strtemp+ls

        return strtemp
   #########################      
    def pr():
        global zhuant
        global sbtime
        global renshu
        global ren_name
        global xj_list
        global qj_list
        global time_data
        global pathfile

        if len(time_data)>5:
            pass
        else:
            tkinter.messagebox.showinfo('警告','请先读取昨日报表')
            return


        nt = datetime.datetime.now()
        qnt = datetime.datetime.now()+datetime.timedelta(days=-1)
        
        document = Document(pathfile)
        #document = Document('demo1.docx')
        document.styles['Normal'].font.name = u'宋体'
        document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        
        pp = document.paragraphs[0].clear()
        document.paragraphs[3].clear()
        pp.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run1 = pp.add_run('报告日期：'+qnt.strftime('%Y{y}%m{m}%d{d}').format(y='年', m='月', d='日')+'0:00至'+nt.strftime('%Y{y}%m{m}%d{d}').format(y='年', m='月', d='日')+'0:00')
        font = run1.font
        font.size = Pt(10)
        font.bold = True

        tables=document.tables[0]
        save_sj=time_cl()
        strtemp=''
        for n in save_sj:
            strtemp+=''.join(n)
            strtemp+='\n'

        tables.cell(2,0).text = strtemp
        tables.cell(2,1).text = text_getstr()
        tables.cell(2,2).text = renyuan()

        document.add_paragraph('        填报人：'+''.join(ren_name[yb_name.get()])+'                                                                                                  审核人：罗肖')

        save_path=tkinter.filedialog.asksaveasfilename()
        strh12.set(save_path)
        if 'docx' not in save_path:
            save_path=save_path+'.docx'
        document.save(save_path)
        tkinter.messagebox.showinfo('数据生成', '日报生成完成！', parent=root)
        
    #>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>
    def text_getstr():#获取文本框内输入内容，返回内容的字符串‘str’
        if textpad.get(1.0,'end')!='':
            words=textpad.get(1.0,'end')
        return words#返回字符串
    
    
    def zuori_data():#读取昨日数据
        global time_data
        global pathfile
        pathfile,time_data=open_word.word_open()
    

#>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>
    root=Tk()
    root.geometry('800x600')
    root.resizable(0,0)
    root.title("北京普瑞浩特能源科技公司潜江压气站")

    frame1=LabelFrame(height=200,width=800,text = '生产动态')
    frame1.grid_propagate(0)
    frame1.grid(column = 0,row = 0,padx = 0,pady = 0)

    h1 = BooleanVar() # 设置选择框对象
    cb1 =Checkbutton(frame1,text='A机组',variable=h1,command=zt_cl).grid(column=0,row = 0)
    strh1=StringVar()
    strh1.set('0')
    strcb1=Entry(frame1,textvariable = strh1,width = 5).grid(column = 1,row = 0)
    
    h2 = BooleanVar() # 设置选择框对象
    cb2 =Checkbutton(frame1,text='B机组',variable=h2,command=zt_cl).grid(column=0,row = 1)
    strh2=StringVar()
    strh2.set('0')
    strcb2=Entry(frame1,textvariable = strh2,width = 5).grid(column = 1,row = 1)

    h3 = BooleanVar() # 设置选择框对象
    cb3 =Checkbutton(frame1,text='C机组',variable=h3,command=zt_cl).grid(column=0,row = 2)
    strh3=StringVar()
    strh3.set('0')
    strcb3=Entry(frame1,textvariable = strh3,width = 5).grid(column = 1,row = 2)

    h4 = BooleanVar() # 设置选择框对象
    cb4 =Checkbutton(frame1,text='D机组',variable=h4,command=zt_cl).grid(column=0,row = 3)
    strh4=StringVar()
    strh4.set('0')
    strcb4=Entry(frame1,textvariable = strh4,width = 5).grid(column = 1,row = 3)

    h5 = BooleanVar() # 设置选择框对象
    cb5 =Checkbutton(frame1,text='1#空压机',variable=h5,command=zt_cl).grid(column=2,row = 0)
    strh5=StringVar()
    strh5.set('0')
    strcb5=Entry(frame1,textvariable = strh5,width = 5).grid(column = 3,row = 0)

    h6 = BooleanVar() # 设置选择框对象
    cb6 = Checkbutton(frame1,text='2#空压机',variable=h6,command=zt_cl).grid(column=2,row = 1)
    strh6=StringVar()
    strh6.set('0')
    strcb6=Entry(frame1,textvariable = strh6,width = 5).grid(column = 3,row = 1)

    h7 = BooleanVar() # 设置选择框对象
    cb7  =Checkbutton(frame1,text='3#空压机',variable=h7,command=zt_cl).grid(column=2,row = 2)
    strh7=StringVar()
    strh7.set('0')
    strcb7=Entry(frame1,textvariable = strh7,width = 5).grid(column = 3,row = 2)

    h8 = BooleanVar() # 设置选择框对象
    cb8  =Checkbutton(frame1,text='1#循环水泵',variable=h8,command=zt_cl).grid(column=4,row = 0)
    strh8=StringVar()
    strh8.set('0')
    strcb8=Entry(frame1,textvariable = strh8,width = 5).grid(column = 5,row = 0)

    h9 = BooleanVar() # 设置选择框对象
    cb9  =Checkbutton(frame1,text='2#循环水泵',variable=h9,command=zt_cl).grid(column=4,row = 1)
    strh9=StringVar()
    strh9.set('0')
    strcb9=Entry(frame1,textvariable = strh9,width = 5).grid(column = 5,row = 1)

    h10 = BooleanVar() # 设置选择框对象
    cb10  =Checkbutton(frame1,text='3#循环水泵',variable=h10,command=zt_cl).grid(column=4,row = 2)
    strh10=StringVar()
    strh10.set('0')
    strcb10=Entry(frame1,textvariable = strh10,width = 5).grid(column = 5,row = 2)

    h11 = BooleanVar() # 设置选择框对象
    cb11  =Checkbutton(frame1,text='4#循环水泵',variable=h11,command=zt_cl).grid(column=4,row = 3)
    strh11=StringVar()
    strh11.set('0')
    strcb11=Entry(frame1,textvariable = strh11,width = 5).grid(column = 5,row = 3)

    label = Label(frame1,text ='文件保存路径:').grid(column = 0,row = 5)
    strh12 = StringVar()
    strcb12 = Entry(frame1,textvariable = strh12,width=40).grid(column = 1,row = 5,columnspan = 5)
    
    button2=Button(frame1,text = '读取昨日数据',command = zuori_data).grid(column = 6,row = 5)
    button=Button(frame1,text = '生成今日数据',command =pr).grid(column =7,row =5)
    button1 = Button(frame1,text = 'QUIT',command = root.quit).grid(column = 8,row = 5)

    labe2 = Label(frame1,text ='A机组流量').grid(column = 0,row = 6)
    llint1=IntVar()
    llint1.set(0)
    llcb2=Entry(frame1,textvariable = llint1,width = 5).grid(column = 1,row = 6)
    
    labe3 = Label(frame1,text ='B机组流量').grid(column = 2,row = 6)
    llint2=IntVar()
    llint2.set(0)
    llcb3=Entry(frame1,textvariable = llint2,width = 5).grid(column = 3,row = 6)
    
    labe4 = Label(frame1,text ='C机组流量').grid(column = 4,row = 6)
    llint3=IntVar()
    llint3.set(0)
    llcb4=Entry(frame1,textvariable = llint3,width = 5).grid(column = 5,row = 6)
    
    labe5 = Label(frame1,text ='D机组流量').grid(column = 6,row = 6)
    llint4=IntVar()
    llint4.set(0)
    llcb5=Entry(frame1,textvariable = llint4,width = 5).grid(column = 7,row = 6)
    #》》》》》》》》》》》》》》》》》》》》》》》》》》》》
    frame2=LabelFrame(height=200,width=800,text = '工作情况')
    frame2.grid_propagate(0)
    frame2.grid(column = 0,row = 1,ipadx = 0,ipady = 0)
    textpad=Text(frame2,undo=True,width = 110,height=13)
    textpad.grid(column = 0, row = 0)
    textpad.insert(END,text_open())
    scroll=Scrollbar(frame2)
    textpad.config(yscrollcommand = scroll.set)
    scroll.config(command = textpad.yview)
    scroll.grid(column = 1, row = 0,stick = 'ns')

    #第三部份人员动迁情况
    frame3=LabelFrame(height=200,width=800,text = '人员及动迁情况')
    frame3.grid_propagate(0)
    frame3.grid(column = 0,row = 2,ipadx = 0,ipady = 0)
    textlable = Label(frame3,text = '潜江站在岗人员：').grid(column = 0,row = 0)
    ren_shu=IntVar()
    entry_renshu=Entry(frame3,textvariable = ren_shu,width = 3).grid(column = 1,row = 0)
    
    rs1 = BooleanVar() # 设置选择框对象
    rscb1 =Checkbutton(frame3,text='罗肖',variable=rs1,command=rs_xz).grid(row = 1,column = 0,stick= 'e')
    rs2 = BooleanVar() # 设置选择框对象
    rscb2 =Checkbutton(frame3,text='郭新兵',variable=rs2,command=rs_xz).grid(row = 1,column = 1)
    rs3 = BooleanVar() # 设置选择框对象
    rscb3 =Checkbutton(frame3,text='赵魏',variable=rs3,command=rs_xz).grid(row = 1,column = 2)
    rs4 = BooleanVar() # 设置选择框对象
    rscb4 =Checkbutton(frame3,text='王灿南',variable=rs4,command=rs_xz).grid(row = 1,column = 3)
    rs5 = BooleanVar() # 设置选择框对象
    rscb5 =Checkbutton(frame3,text='聂向东',variable=rs5,command=rs_xz).grid(row = 1,column = 4)
    rs6 = BooleanVar()
    rscb6 = Checkbutton(frame3,text ='谢强',variable=rs6,command = rs_xz).grid(row =1,column = 5)

    textlable2=Label(frame3,text = '潜江站夜班人员：').grid(row = 2,column = 0)
    yb_name=IntVar()
    yb_name.set(0)
    yb_column=0
    for n in ren_name:
        yb=Radiobutton(frame3,text = n,variable = yb_name,value =yb_column)
        yb.grid(row = 2,column = yb_column+1)
        yb_column+=1

    textlable3=Label(frame3,text = '潜江站休假人员：').grid(row = 3,column = 0)
    xj1 = BooleanVar() # 设置选择框对象
    xjcb1 =Checkbutton(frame3,text='罗肖',variable=xj1,command=xj_xz).grid(row = 3,column = 1,stick= 'e')
    xj2 = BooleanVar() # 设置选择框对象
    xjcb2 =Checkbutton(frame3,text='郭新兵',variable=xj2,command=xj_xz).grid(row = 3,column = 2)
    xj3 = BooleanVar() # 设置选择框对象
    xjcb3 =Checkbutton(frame3,text='赵魏',variable=xj3,command=xj_xz).grid(row = 3,column = 3)
    xj4 = BooleanVar() # 设置选择框对象
    xjcb4 =Checkbutton(frame3,text='王灿南',variable=xj4,command=xj_xz).grid(row = 3,column = 4)
    xj5 = BooleanVar() # 设置选择框对象
    xjcb5 =Checkbutton(frame3,text='聂向东',variable=xj5,command=xj_xz).grid(row = 3,column = 5)
    xj6 = BooleanVar()
    xjcb6 = Checkbutton(frame3,text ='谢强',variable=xj6,command = xj_xz).grid(row =3,column = 6)
    
    textlable4=Label(frame3,text = '潜江站请假人员：').grid(row = 4,column = 0)
    qj1 = BooleanVar() # 设置选择框对象
    qjcb1 =Checkbutton(frame3,text='罗肖',variable=qj1,command=qj_xz).grid(row = 4,column = 1,stick= 'e')
    qj2 = BooleanVar() # 设置选择框对象
    qjcb2 =Checkbutton(frame3,text='郭新兵',variable=qj2,command=qj_xz).grid(row = 4,column = 2)
    qj3 = BooleanVar() # 设置选择框对象
    qjcb3 =Checkbutton(frame3,text='赵魏',variable=qj3,command=qj_xz).grid(row = 4,column = 3)
    qj4 = BooleanVar() # 设置选择框对象
    qjcb4 =Checkbutton(frame3,text='王灿南',variable=qj4,command=qj_xz).grid(row = 4,column = 4)
    qj5 = BooleanVar() # 设置选择框对象
    qjcb5 =Checkbutton(frame3,text='聂向东',variable=qj5,command=qj_xz).grid(row = 4,column = 5)
    qj6 = BooleanVar()
    qjcb6 = Checkbutton(frame3,text ='谢强',variable=qj6,command = qj_xz).grid(row =4,column = 6)
    
    bltext()
    root.mainloop()
##################################
def str_time_cl(str_sj):#洗出时间
    temp_s=0
    temp1=''
    for n in str_sj:
        if  '0'<=n<='9':
            temp_s=temp_s*10+int(n)
        else:
            break
    if '.' in str_sj:
        temp1=str_sj[-2:]
        miao=temp_s*60+int(temp1)
    else:
        miao=temp_s*60
    return miao
#########################
def time_cl():#输出
    global sbtime
    global zhuant
    global sb_name1
    global time_data
    global liulian
    time_end=[]
    data_end=[]
    str_temp=''
    time_end.append(["潜江站："])
    for n in range(len(sb_name1)):
        if zhuant[n]==0:
            xiaoshi,fenzhong=time_str(str_time_cl(sbtime[n])+str_time_cl(time_data[n]))
            if fenzhong=='0':
                fenzhong='00'
            data_end.append(xiaoshi+'.'+fenzhong)
            str_temp=sb_name1[n]+'设备未运行，累计运行'+xiaoshi+'小时。'
        else:
            xiaoshi,fenzhong=time_str(str_time_cl(sbtime[n])+str_time_cl(time_data[n]))
            if fenzhong=='0':
                fenzhong='00'
            data_end.append(xiaoshi+'.'+fenzhong)
            str_temp=sb_name1[n]+'设备当日运行'+sbtime[n]+'小时，累计运行'+xiaoshi+'小时。'
       
        time_end.append([str_temp])
    time_end.append(["注：当日输气量"+str(liulian)+"万N㎥；"])
    return time_end
##########################
def time_str(shijian):#切换回小时分钟
    xiaoshi=shijian//60
    fenzhong=shijian%60
    return str(xiaoshi),str(fenzhong)
################################
def text_open():#默认文本内容读入 ，如果默认文本不存在者创建一个  
    try:
        os.path.getsize('moren.txt')
    except:
        po=open('moren.txt',"w")
        po.write('\n')
        po.close()
    op=open('moren.txt',"r")
    text_o=[]
    ls=op.read()
    for n in ls:
        text_o.append(n)
    return ''.join(text_o)
############################
if __name__=='__main__':
    pathfile=''
    ren_name=["罗肖","郭新兵","赵魏","王灿南","聂向东","谢强"]
    xj_list=[0,0,0,0,0,0]#休假
    qj_list=[0,0,0,0,0,0]#请假
    renshu=[0,0,0,0,0,0]#在岗人数
    sb_name1=['A机组','B机组','C机组','D机组','1#空压机','2#空压机','3#空压机','1#循坏水泵','2#循坏水泵','3#循坏水泵','4#循坏水泵']
    zhuant=[0,0,0,0,0,0,0,0,0,0,0]
    sbtime=['0','0','0','0','0','0','0','0','0','0','0']
    time_data=()
    liulian=0#机组流量
    mywin()